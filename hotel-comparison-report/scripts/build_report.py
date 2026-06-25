#!/usr/bin/env python3
"""Build a JC Room Blocks Hotel Comparison Report (.docx) from a JSON data file.

Every formatting constant in this script was extracted from the approved
Hepburn/Donaldson reference report. Do NOT change formatting values here to
suit one report -- the whole point of this script is that every report comes
out identical. Only the JSON data changes between reports.

Usage:
    python3 build_report.py <data.json> <output.docx>

After writing the file, the script re-opens it and audits the paragraph
formatting (spacer sizes, table widths, spacing before/after) and prints a
verification report. A non-zero exit code means the audit failed.
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ASSETS = Path(__file__).resolve().parent.parent / "assets"
JC_LOGO = ASSETS / "jc_logo.png"
TA_LOGO = ASSETS / "tripadvisor_logo.jpg"

# ---- formatting constants (from the Hepburn/Donaldson reference) ----
BODY_FONT = "Palatino Linotype"
BODY_SIZE = Pt(10)
LINE_115 = 1.15                 # body line spacing
BULLET_SIZE = Pt(9)
NOTE_SIZE = Pt(8)
RANKING_FONT = "Arial"
RANKING_SIZE = Pt(10.5)
TABLE_FONT = "Calibri"
TABLE_SIZE = Pt(11)
HYPERLINK_BLUE = "0000FF"
SHADE_GRAY = "D9D9D9"

JC_LOGO_W = Emu(2468880)        # 2.70 in
JC_LOGO_H = Emu(1101339)
TA_LOGO_W = Emu(1417320)        # 1.55 in
TA_LOGO_H = Emu(250576)

PAGE_USABLE_DXA = 10656         # 8.5in page - 2 x 0.55in margins, in twips
OFFERED_COL_DXA = 1728
RACK_COL_DXA = 2232
ROOM_COL_MIN_DXA = 2833
ROOM_COL_MAX_DXA = PAGE_USABLE_DXA - OFFERED_COL_DXA - RACK_COL_DXA - 200

BULLET_INDENT = 950             # twips; hanging 346
SUB_INDENT = 1267

DEFAULT_RATE_NOTES = [
    "*Rates are per night and are based on s/d occupancy",
    "*Rates do not include taxes and fees",
]
DEFAULT_CONTACT = ("JC Room Blocks, P.O. Box 2661, Canyon Country, CA 91386 / "
                   "818-406-2351 / Info@JCRoomBlocks.com / www.JCRoomBlocks.com")
CUTOFF_PREFIX = "Cut-off Date: (last date for guests to book at the group rate): "


# ---------------------------------------------------------------- helpers

def set_par(p, *, line=None, align=None, after_pt=0, before_pt=0):
    """Explicitly pin paragraph spacing so Word defaults can never leak in."""
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    if line is not None:
        pf.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def set_mark_size(p, size_pt):
    """Set the paragraph-mark font size (what controls an empty line's height)."""
    pPr = p._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)


def spacer(doc, size_pt):
    """One intentional blank spacer paragraph: exact size, 0 before/after, single."""
    p = doc.add_paragraph()
    set_par(p, line=1.0)
    set_mark_size(p, size_pt)
    return p


def run(p, text, *, size=None, bold=False, font=None, color=None, underline=False):
    r = p.add_run(text)
    if size is not None:
        r.font.size = size
    if bold:
        r.bold = True
    if font:
        r.font.name = font
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if underline:
        r.font.underline = True
    return r


def hyperlink(p, text, url, *, size=None, bold=False, font=None):
    """Add a blue underlined hyperlink run to a paragraph."""
    part = p.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if font:
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), font)
        rf.set(qn("w:hAnsi"), font)
        rf.set(qn("w:cs"), font)
        rPr.append(rf)
    if bold:
        rPr.append(OxmlElement("w:b"))
    c = OxmlElement("w:color")
    c.set(qn("w:val"), HYPERLINK_BLUE)
    rPr.append(c)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))
        rPr.append(sz)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    h.append(r)
    p._p.append(h)
    return h


def bullet(doc, text, sub=False):
    p = doc.add_paragraph()
    set_par(p, line=LINE_115)
    ind = p.paragraph_format
    ind.left_indent = Pt((SUB_INDENT if sub else BULLET_INDENT) / 20)
    pPr = p._p.get_or_add_pPr()
    indEl = pPr.find(qn("w:ind"))
    indEl.set(qn("w:hanging"), "346")
    run(p, "o" if sub else "•", size=BULLET_SIZE)
    tab_r = p.add_run()
    tab_r.font.size = BULLET_SIZE
    tab_r._r.append(OxmlElement("w:tab"))
    run(p, text, size=BULLET_SIZE)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    set_par(p, line=LINE_115)
    run(p, text, bold=True)
    return p


def divider(doc):
    """Horizontal divider line matching the reference report."""
    p = doc.add_paragraph()
    set_par(p, line=1.0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "thinThickThinMediumGap")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_cell(cell, text, *, bold=False, center=False, shade=None, valign):
    cell.vertical_alignment = valign
    tcPr = cell._tc.get_or_add_tcPr()
    # thin black borders
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "000000")
        borders.append(e)
    tcPr.append(borders)
    # compact cell margins
    mar = OxmlElement("w:tcMar")
    for edge, w in (("top", "30"), ("left", "60"), ("bottom", "30"), ("right", "60")):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), w)
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tcPr.append(mar)
    if shade:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), shade)
        tcPr.append(shd)
    p = cell.paragraphs[0]
    set_par(p, line=1.0,
            align=WD_ALIGN_PARAGRAPH.CENTER if center else None)
    run(p, text, size=TABLE_SIZE, bold=bold, font=TABLE_FONT)


def room_col_width(rates):
    """Size the Room Type column to its longest entry, capped to the page."""
    longest = max((len(r["room_type"]) for r in rates), default=20)
    w = 115 * longest + 220
    return max(ROOM_COL_MIN_DXA, min(w, ROOM_COL_MAX_DXA))


def rate_table(doc, rates):
    table = doc.add_table(rows=len(rates) + 1, cols=3)
    room_w = room_col_width(rates)
    widths = (room_w, OFFERED_COL_DXA, RACK_COL_DXA)

    tbl = table._tbl
    tblPr = tbl.tblPr
    # fixed layout, centered, explicit width
    for tag, attrs in (("w:tblW", {"w:w": str(sum(widths)), "w:type": "dxa"}),
                       ("w:jc", {"w:val": "center"}),
                       ("w:tblLayout", {"w:type": "fixed"})):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        tblPr.append(el)
    grid = tbl.find(qn("w:tblGrid"))
    for col, w in zip(grid.findall(qn("w:gridCol")), widths):
        col.set(qn("w:w"), str(w))

    hdr = table.rows[0]
    set_cell(hdr.cells[0], "Room Type", bold=True, valign=WD_ALIGN_VERTICAL.BOTTOM)
    set_cell(hdr.cells[1], "Offered Rate", bold=True, center=True,
             shade=SHADE_GRAY, valign=WD_ALIGN_VERTICAL.BOTTOM)
    set_cell(hdr.cells[2], 'Regular "Rack" Rate', bold=True, center=True,
             valign=WD_ALIGN_VERTICAL.BOTTOM)
    for i, rate in enumerate(rates, start=1):
        row = table.rows[i]
        set_cell(row.cells[0], rate["room_type"], valign=WD_ALIGN_VERTICAL.CENTER)
        set_cell(row.cells[1], rate["offered_rate"], bold=True, center=True,
                 shade=SHADE_GRAY, valign=WD_ALIGN_VERTICAL.CENTER)
        set_cell(row.cells[2], rate["rack_rate"], center=True,
                 valign=WD_ALIGN_VERTICAL.CENTER)
    # explicit cell widths so the fixed layout is honored everywhere
    for r in table.rows:
        for c, w in zip(r.cells, widths):
            tcW = c._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                c._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
    return table


# ---------------------------------------------------------------- document

def add_hotel(doc, hotel, venue, needs_verification):
    # Hotel name (hyperlinked to the hotel website when a URL is provided)
    p = doc.add_paragraph()
    set_par(p, line=LINE_115)
    if hotel.get("website_url"):
        hyperlink(p, hotel["name"], hotel["website_url"], size=Pt(12), bold=True)
    else:
        run(p, hotel["name"], size=Pt(12), bold=True)
    for line in (hotel["address"], hotel["city_state_zip"]):
        set_par(doc.add_paragraph(), line=LINE_115)
        run(doc.paragraphs[-1], line)

    spacer(doc, 6)

    # TripAdvisor logo + ranking line
    p = doc.add_paragraph()
    set_par(p, line=LINE_115)
    p.add_run().add_picture(str(TA_LOGO), width=TA_LOGO_W, height=TA_LOGO_H)
    p = doc.add_paragraph()
    set_par(p, line=LINE_115)
    ta = hotel.get("tripadvisor") or {}
    if ta.get("ranking"):
        text = ta["ranking"]
    else:
        text = "TripAdvisor rating not verified"
        if ta.get("note"):
            text += f" - {ta['note']}"
        needs_verification.append(f"{hotel['name']}: TripAdvisor ranking")
    if ta.get("url"):
        hyperlink(p, text, ta["url"], size=RANKING_SIZE, font=RANKING_FONT)
    else:
        run(p, text, size=RANKING_SIZE, font=RANKING_FONT)
        if ta.get("ranking"):
            needs_verification.append(f"{hotel['name']}: TripAdvisor URL missing")

    spacer(doc, 6)

    heading(doc, "Special Negotiated Group Rates:")
    rate_table(doc, hotel["rates"])
    for note in hotel.get("rate_notes", DEFAULT_RATE_NOTES):
        p = doc.add_paragraph()
        set_par(p, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        run(p, note, size=NOTE_SIZE)

    heading(doc, "Hotel Features:")
    for item in hotel["features"]:
        bullet(doc, item.lstrip("> "), sub=item.startswith(">"))
    spacer(doc, 6)

    heading(doc, "Concessions:")
    for item in hotel["concessions"]:
        bullet(doc, item.lstrip("> "), sub=item.startswith(">"))
    spacer(doc, 6)

    heading(doc, "Contracting Option:")
    for item in hotel["contracting_options"]:
        bullet(doc, item.lstrip("> "), sub=item.startswith(">"))
    spacer(doc, 6)

    p = doc.add_paragraph()
    set_par(p, line=LINE_115)
    run(p, f"Distance from {venue}: {hotel['distance_from_venue']}", bold=True)
    p = doc.add_paragraph()
    set_par(p, line=LINE_115)
    run(p, CUTOFF_PREFIX + hotel["cutoff_date"], bold=True)


def build(data, out_path):
    doc = Document()

    # page + base style
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Emu(7772400), Emu(10058400)  # 8.5 x 11
    sec.left_margin = sec.right_margin = Emu(503555)               # 0.55 in
    sec.top_margin = sec.bottom_margin = Emu(320040)               # 0.35 in
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    # title block
    p = doc.add_paragraph()
    set_par(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(str(JC_LOGO), width=JC_LOGO_W, height=JC_LOGO_H)
    p = doc.add_paragraph()
    set_par(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Hotel Comparison Report", size=Pt(14), bold=True, underline=True)
    p = doc.add_paragraph()
    set_par(p, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, f"Prepared for {data['prepared_for']}", size=Pt(14))
    p = doc.add_paragraph()
    set_par(p, align=WD_ALIGN_PARAGRAPH.CENTER, after_pt=14)
    run(p, data["event_dates"], size=Pt(12))

    needs_verification = []
    venue = data["venue_name"]
    for i, hotel in enumerate(data["hotels"]):
        add_hotel(doc, hotel, venue, needs_verification)
        # between hotels and before the contact line: 10pt / divider / 10pt
        spacer(doc, 10)
        divider(doc)
        spacer(doc, 10)

    p = doc.add_paragraph()
    set_par(p, line=1.0)
    run(p, data.get("contact_line", DEFAULT_CONTACT), size=Pt(9))

    doc.save(out_path)
    return needs_verification


# ---------------------------------------------------------------- audit

def audit(path):
    """Re-open the saved file and verify the formatting contract."""
    doc = Document(path)
    problems = []
    empty_sizes = []
    for p in doc.paragraphs:
        pf = p.paragraph_format
        if pf.space_before is not None and pf.space_before.pt not in (0, 14):
            problems.append(f"space_before={pf.space_before.pt} on {p.text[:40]!r}")
        is_empty = not p.text.strip() and "blip" not in p._p.xml
        if is_empty and "pBdr" not in p._p.xml:
            pPr = p._p.find(qn("w:pPr"))
            sz = None
            if pPr is not None:
                rPr = pPr.find(qn("w:rPr"))
                if rPr is not None:
                    szEl = rPr.find(qn("w:sz"))
                    if szEl is not None:
                        sz = int(szEl.get(qn("w:val"))) / 2
            empty_sizes.append(sz)
            if sz not in (6.0, 10.0):
                problems.append(f"blank paragraph with size {sz} (must be 6 or 10)")
    for i, t in enumerate(doc.tables):
        grid = t._tbl.find(qn("w:tblGrid"))
        total = sum(int(c.get(qn("w:w"))) for c in grid.findall(qn("w:gridCol")))
        if total > PAGE_USABLE_DXA:
            problems.append(f"table {i} width {total} dxa exceeds page ({PAGE_USABLE_DXA})")
    n6 = empty_sizes.count(6.0)
    n10 = empty_sizes.count(10.0)
    print(f"Audit: {n6} six-pt spacers, {n10} ten-pt spacers, "
          f"{len(doc.tables)} rate tables, all within page width."
          if not problems else "Audit FAILED:")
    for pr in problems:
        print("  -", pr)
    return not problems


def main():
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    data = json.loads(Path(sys.argv[1]).read_text())
    out = sys.argv[2]
    needs_verification = build(data, out)
    ok = audit(out)
    print(f"Saved: {out}")
    if needs_verification:
        print("NEEDS VERIFICATION:")
        for item in needs_verification:
            print("  -", item)
    sys.exit(0 if ok else 1)


if __name__ == "__main__":
    main()
