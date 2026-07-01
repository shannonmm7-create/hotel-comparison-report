"""Render the Hotel Comparison Report ``.docx`` from validated data.

The whole engine is docxtpl (Jinja2 over Word XML). The only non-obvious parts —
currency formatting, per-hotel RichText hyperlinks, and why ``autoescape`` is
mandatory — are explained inline and in ``docs/docx-templating-notes.md``.
"""

from __future__ import annotations

import io
import re
from collections.abc import Mapping
from decimal import ROUND_HALF_UP, Decimal
from importlib.resources import files
from pathlib import Path
from typing import Any, TypedDict

from docx import Document
from docx.oxml.ns import qn
from docxtpl import DocxTemplate, RichText  # type: ignore[import-untyped]  # docxtpl ships no py.typed
from jinja2 import Environment

from hotel_report._validate import validated
from hotel_report.models import Report

DEFAULT_TEMPLATE = files("hotel_report").joinpath("templates/hotel_comparison_report.docx")


class RenderSummary(TypedDict):
    """Summary returned by :func:`render`."""

    output: str
    hotels: int
    rooms: int


# Matches the approved hyperlink look captured from the source template.
_LINK_BLUE = "0000FF"
_NAME_SIZE = 24  # half-points -> 12pt
_TA_SIZE = 21  # half-points -> 10.5pt

# A whole Jinja token ({{ }}, {% %}, {# #}); used to detect unrendered residue.
_TOKEN_RE = re.compile(r"\{\{.*?\}\}|\{%.*?%\}|\{#.*?#\}", re.S)


class RenderError(RuntimeError):
    """Raised when a rendered document still contains unrendered template tokens."""


@validated
def usd(value: float | int | str) -> str:
    """Format a rate as US currency. Numbers get a ``$`` and thousands
    separators (279 -> $279, 1250 -> $1,250); non-integers round half-up via
    ``Decimal`` so binary-float amounts like 2.675 render correctly ($2.68).
    Strings pass through unchanged, so the caller controls edge cases
    ('Waived', '$305 + tax')."""
    if isinstance(value, (int, float)):
        amount = Decimal(str(value))  # str() avoids the binary-float repr
        if amount == amount.to_integral_value():
            return f"${amount:,.0f}"
        return f"${amount.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP):,.2f}"
    return str(value)


@validated
def _jinja_env() -> Environment:
    """Build the Jinja environment (autoescape on) with the ``usd`` filter registered."""
    # autoescape=True is REQUIRED: data may contain & < > which would otherwise
    # produce invalid document XML. docxtpl's RichText is autoescape-safe, so
    # the hyperlinks below still render as real links.
    env = Environment(autoescape=True)
    env.filters["usd"] = usd
    return env


@validated
def _build_context(tpl: DocxTemplate, report: Report) -> dict[str, Any]:
    """Turn a :class:`Report` into the docxtpl render context, adding a RichText
    hyperlink for each hotel's name and TripAdvisor line (see ``docs/`` for why
    the links must be RichText)."""
    ctx: dict[str, Any] = report.model_dump()
    for hotel_ctx, hotel in zip(ctx["hotels"], report.hotels, strict=True):
        # hotel name -> RichText hyperlink (approved look: bold, underline, blue, 12pt)
        url = hotel.website_url
        name = RichText()
        name.add(
            hotel.name,
            url_id=tpl.build_url_id(str(url)) if url else None,
            bold=True,
            underline=bool(url),
            color=_LINK_BLUE if url else "000000",
            size=_NAME_SIZE,
        )
        hotel_ctx["name_link"] = name
        # tripadvisor -> RichText hyperlink (underline, blue, 10.5pt, Arial); when
        # omitted, say so explicitly rather than leaving a silent blank line.
        ta = hotel.tripadvisor
        ta_url = ta.url if ta else None
        ta_link = RichText()
        ta_link.add(
            ta.text if ta else "TripAdvisor rating not verified",
            url_id=tpl.build_url_id(str(ta_url)) if ta_url else None,
            underline=bool(ta_url),
            color=_LINK_BLUE if ta_url else "000000",
            size=_TA_SIZE,
            font="Arial",
        )
        hotel_ctx["tripadvisor_link"] = ta_link
        # sub-bullets: a leading "> " marks an indented sub-item under the previous one.
        for key in ("features", "concessions", "contracting_options"):
            hotel_ctx[key] = [_format_bullet(item) for item in hotel_ctx[key]]
    return ctx


@validated
def _format_bullet(item: str) -> str:
    """Convert a ``"> "``-prefixed item into an indented sub-bullet (strip the
    marker, add a leading tab); return other items unchanged."""
    return "\t" + item[2:] if item.startswith("> ") else item


@validated
def _part_elements(doc: Any) -> list[Any]:
    """The body plus every header/footer element — docxtpl renders all of them,
    so residue detection must look at all of them too. Duplicates are harmless
    (residue is computed as a set intersection)."""
    parts = [doc.element.body]
    for sec in doc.sections:
        for hf in (
            sec.header,
            sec.first_page_header,
            sec.even_page_header,
            sec.footer,
            sec.first_page_footer,
            sec.even_page_footer,
        ):
            parts.append(hf._element)  # pylint: disable=protected-access  # python-docx internal
    return parts


@validated
def _jinja_tokens(doc: Any) -> list[str]:
    """Every Jinja token in the document, coalescing runs per paragraph (Word
    can split a token across runs) across body, headers and footers."""
    toks: list[str] = []
    for part in _part_elements(doc):
        for p in part.iter(qn("w:p")):
            line = "".join((t.text or "") for t in p.iter(qn("w:t")))
            toks.extend(_TOKEN_RE.findall(line))
    return toks


@validated
def find_residue(path: str | Path) -> list[str]:
    """Jinja tokens surviving anywhere in a rendered document (body/headers/footers)."""
    return _jinja_tokens(Document(str(path)))


@validated
def _control_tokens(tokens: list[str]) -> set[str]:
    """The control tags ({% %} / {# #}) among ``tokens``. Expression tags ({{ }})
    are excluded from residue detection: a surviving ``{{ }}`` is almost always
    data that happens to contain braces, whereas a surviving control tag is a
    real template bug."""
    return {t for t in tokens if t.startswith(("{%", "{#"))}


@validated
def render(
    data: Report | Mapping[str, Any],
    output_path: str | Path,
    template: str | Path | None = None,
) -> RenderSummary:
    """Fill the template from ``data`` and write ``output_path``.

    ``data`` may be a :class:`~hotel_report.models.Report` or any mapping that
    parses into one (a :class:`pydantic.ValidationError` is raised otherwise).
    Raises :class:`RenderError` only if a *control tag* from the template fails to
    render (a template bug); ``{{ }}`` expressions and data that merely contains
    braces are never flagged."""
    report = data if isinstance(data, Report) else Report.model_validate(data)
    if template is None:
        raw = DEFAULT_TEMPLATE.read_bytes()  # bytes so it works from a zipped install too
        template_ctrl = _control_tokens(_jinja_tokens(Document(io.BytesIO(raw))))
        tpl = DocxTemplate(io.BytesIO(raw))
    else:
        template_ctrl = _control_tokens(_jinja_tokens(Document(str(template))))
        tpl = DocxTemplate(str(template))
    tpl.render(_build_context(tpl, report), jinja_env=_jinja_env())
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(output_path))
    # scan the saved (rendered) document for surviving control tags
    residue = sorted(_control_tokens(find_residue(output_path)) & template_ctrl)
    if residue:
        raise RenderError(f"unrendered template tokens remain: {residue}")
    return RenderSummary(
        output=str(output_path),
        hotels=len(report.hotels),
        rooms=sum(len(h.rooms) for h in report.hotels),
    )
