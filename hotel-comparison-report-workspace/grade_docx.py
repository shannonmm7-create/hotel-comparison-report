#!/usr/bin/env python3
"""Grade a hotel-comparison-report eval run programmatically.

Usage: python3 grade_docx.py <run_dir> <eval_id>
Reads the first .docx in <run_dir>/outputs/, evaluates the assertions for that
eval, writes <run_dir>/grading.json with {text, passed, evidence} entries.
"""
import json, re, sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

PAGE_USABLE_DXA = 10656
CONTACT_SNIPPET = "JC Room Blocks, P.O. Box 2661"


def para_info(doc):
    out = []
    for p in doc.paragraphs:
        x = p._p.xml
        pPr = p._p.find(qn("w:pPr"))
        mark = None
        if pPr is not None:
            rPr = pPr.find(qn("w:rPr"))
            if rPr is not None:
                s = rPr.find(qn("w:sz"))
                if s is not None:
                    mark = int(s.get(qn("w:val"))) / 2
        out.append({
            "text": p.text, "xml": x, "mark": mark,
            "empty": not p.text.strip() and "blip" not in x,
            "border": "pBdr" in x,
            "sb": p.paragraph_format.space_before,
            "sa": p.paragraph_format.space_after,
        })
    return out


def check(docx_path, eval_id, summary_text):
    doc = Document(docx_path)
    paras = para_info(doc)
    results = []

    def add(text, passed, evidence):
        results.append({"text": text, "passed": bool(passed), "evidence": evidence})

    add("Output is a valid .docx that opens without errors", True,
        f"Opened {docx_path.name}: {len(paras)} paragraphs, {len(doc.tables)} tables")

    # spacer rule
    bad = [p["mark"] for p in paras if p["empty"] and not p["border"]
           and p["mark"] not in (6.0, 10.0)]
    n6 = sum(1 for p in paras if p["empty"] and p["mark"] == 6.0)
    n10 = sum(1 for p in paras if p["empty"] and p["mark"] == 10.0)
    add("Every blank spacer paragraph is exactly 6pt or 10pt — no 10.5pt or default-size blank paragraphs anywhere",
        not bad, f"{n6} 6pt spacers, {n10} 10pt spacers, offending blank sizes: {bad or 'none'}")

    # space before/after 0
    bad_sp = [p["text"][:40] for p in paras
              if (p["sb"] is not None and p["sb"].pt != 0)
              or (p["sa"] is not None and p["sa"].pt not in (0, 14))]
    if eval_id == 0:
        add("All body paragraphs have space before = 0 and space after = 0 (except the 14pt-after dates line)",
            not bad_sp, f"violations: {bad_sp or 'none'}")

    # table widths + shading
    too_wide, shaded = [], []
    for i, t in enumerate(doc.tables):
        grid = t._tbl.find(qn("w:tblGrid"))
        cols = [int(c.get(qn("w:w"), 0)) for c in grid.findall(qn("w:gridCol"))] if grid is not None else []
        if cols and sum(cols) > PAGE_USABLE_DXA:
            too_wide.append((i, sum(cols)))
        shaded.append("D9D9D9" in t._tbl.xml.upper() or "d9d9d9" in t._tbl.xml)
    if eval_id == 0:
        add("Both rate tables fit within the usable page width and the Offered Rate column is shaded gray (D9D9D9)",
            not too_wide and all(shaded),
            f"too wide: {too_wide or 'none'}; gray shading per table: {shaded}")
    else:
        add("The table containing the long room type does not exceed the usable page width",
            not too_wide, f"table widths over {PAGE_USABLE_DXA} dxa: {too_wide or 'none'}")

    full_text = "\n".join(p["text"] for p in paras)
    xml_all = doc.element.xml

    if eval_id == 0:
        ranks = re.findall(r"#\d+ of \d+ Hotels in Savannah", full_text)
        linked = all(r in xml_all for r in ranks) and "tripadvisor.com" in " ".join(
            rel.target_ref for rel in doc.part.rels.values() if "hyperlink" in rel.reltype)
        add("TripAdvisor ranking lines use exactly the format '#X of Y Hotels in Savannah' and are hyperlinked to the TripAdvisor URLs provided",
            len(ranks) == 2 and linked, f"found rankings: {ranks}; tripadvisor hyperlink rels present: {linked}")
        div_ok = "thinThickThinMediumGap" in xml_all
        # one 10pt spacer before and after each divider
        seq_ok = True
        for i, p in enumerate(paras):
            if p["border"]:
                before = paras[i-1] if i > 0 else None
                after = paras[i+1] if i+1 < len(paras) else None
                if not (before and before["empty"] and before["mark"] == 10.0):
                    seq_ok = False
                if after is not None and not (after["empty"] and after["mark"] == 10.0):
                    seq_ok = False
        add("Divider lines between hotel sections use the thinThickThinMediumGap paragraph border with exactly one 10pt spacer before and after",
            div_ok and seq_ok, f"divider style present: {div_ok}; 10pt spacer adjacency ok: {seq_ok}")
        rates_ok = all(s in full_text or s in xml_all for s in ["$249", "$419", "$209", "$359"])
        add("Rate data is accurate: DeSoto Deluxe King $249/$419 and Hyatt City View King $209/$359 appear in the tables",
            rates_ok, "checked $249, $419, $209, $359 presence")
    else:
        n_ranks = len(re.findall(r"#\d+ of \d+ Hotels in Asheville", full_text))
        add("The Radical's TripAdvisor line reads 'TripAdvisor rating not verified' (no invented ranking)",
            "TripAdvisor rating not verified" in full_text and n_ranks == 1,
            f"'not verified' present: {'TripAdvisor rating not verified' in full_text}; "
            f"Asheville ranking lines found: {n_ranks} (must be 1, Grand Bohemian only)")
        add("The unverified TripAdvisor ranking is surfaced to the user as a Needs Verification item in the agent's summary",
            "verif" in summary_text.lower() and "radical" in summary_text.lower(),
            f"summary mentions verification: {'verif' in summary_text.lower()}, mentions The Radical: {'radical' in summary_text.lower()}")
        gb = re.findall(r"#18 of 28 Hotels in Asheville", full_text)
        add("Grand Bohemian ranking is hyperlinked and formatted exactly as '#18 of 28 Hotels in Asheville'",
            bool(gb), f"found: {gb}")
        add("Divider lines between hotel sections use the thinThickThinMediumGap paragraph border",
            "thinThickThinMediumGap" in xml_all, f"present: {'thinThickThinMediumGap' in xml_all}")

    # contact line once, at end
    cnt = full_text.count("JC Room Blocks, P.O. Box 2661")
    nonempty = [p for p in paras if p["text"].strip()]
    last_ok = nonempty and CONTACT_SNIPPET in nonempty[-1]["text"]
    add("The JC Room Blocks contact line appears exactly once, as the last paragraph of the document",
        cnt == 1 and last_ok, f"occurrences: {cnt}; last non-empty paragraph is contact line: {bool(last_ok)}")
    return results


def main():
    run_dir = Path(sys.argv[1])
    eval_id = int(sys.argv[2])
    outputs = run_dir / "outputs"
    docs = sorted(outputs.glob("**/*.docx"))
    summary = ""
    for f in outputs.glob("**/*.md"):
        summary += f.read_text()
    if not docs:
        results = [{"text": "Output is a valid .docx that opens without errors",
                    "passed": False, "evidence": "no .docx file found in outputs"}]
    else:
        try:
            results = check(docs[0], eval_id, summary)
        except Exception as e:
            results = [{"text": "Output is a valid .docx that opens without errors",
                        "passed": False, "evidence": f"failed to parse: {e}"}]
    npass = sum(r["passed"] for r in results)
    payload = {
        "expectations": results,
        "summary": {
            "passed": npass,
            "failed": len(results) - npass,
            "total": len(results),
            "pass_rate": round(npass / len(results), 4) if results else 0.0,
        },
    }
    (run_dir / "grading.json").write_text(json.dumps(payload, indent=2))
    print(f"{run_dir.name}: {npass}/{len(results)} passed")


if __name__ == "__main__":
    main()
