#!/usr/bin/env python3
"""Build hotel comparison report for the Birch/Kowalski wedding."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Birch-Kowalski_Wedding_Hotel_Comparison_Report.docx")

ACCENT = "5B3A66"  # plum
LIGHT = "EFE6F2"
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_RGB = RGBColor(0x5B, 0x3A, 0x66)

doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

def set_cell_bg(cell, color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hl.append(run)
    paragraph._p.append(hl)

def heading(text, size=14, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = ACCENT_RGB
    return p

# ---------- Title block ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Hotel Comparison Report")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = ACCENT_RGB

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Birch / Kowalski Wedding")
r.font.size = Pt(15)
r.bold = True

details = doc.add_paragraph()
details.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = details.add_run("May 14–16, 2027  |  Venue: The Biltmore Estate, Asheville, NC")
r.font.size = Pt(12)

prep = doc.add_paragraph()
prep.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = prep.add_run("Prepared June 11, 2026")
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

# horizontal rule via bottom border
pr = prep._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8")
bottom.set(qn("w:space"), "6"); bottom.set(qn("w:color"), ACCENT)
pbdr.append(bottom)
pr.append(pbdr)

# ---------- Overview ----------
heading("Overview", size=15)
p = doc.add_paragraph(
    "This report compares two hotel options for the Birch/Kowalski wedding room block over the "
    "weekend of May 14–16, 2027. The ceremony and reception will take place at the Biltmore "
    "Estate in Asheville, North Carolina. Both properties were evaluated on guest room rates, "
    "amenities, concessions, contract terms, distance to the venue, and booking deadlines. "
    "Key differences and a recommendation are summarized at the end of this report."
)
p.paragraph_format.space_after = Pt(8)

# ---------- Side-by-side comparison table ----------
heading("Side-by-Side Comparison", size=15)

rows = [
    ("Address", "11 Boston Way, Asheville, NC 28803", "95 Roberts St, Asheville, NC 28801"),
    ("Website", "WEBSITE1", "WEBSITE2"),
    ("TripAdvisor Ranking", "#18 of 28 hotels in Asheville (TALINK)",
     "No TripAdvisor ranking available (newer property)"),
    ("Group Rates",
     "Premier Two Queen Mountain View Suite with Balcony and Kitchenette: $389/night (reg. $599)\n"
     "Classic King: $339/night (reg. $529)",
     "King Studio: $279/night (reg. $409)\n"
     "Double Queen Studio: $289/night (reg. $419)"),
    ("Dining & Amenities",
     "On-site restaurant & bar (Red Stag Grill); art gallery; fitness center",
     "Rooftop restaurant & bar (Golden Hour); on-site café; fitness center"),
    ("Parking", "Valet parking, $35/night (subject to change)", "Self-parking, $25/night (subject to change)"),
    ("Concessions",
     "Complimentary Wi-Fi; complimentary gift bag distribution at front desk (non-personalized)",
     "Complimentary Wi-Fi; complimentary welcome drink at check-in"),
    ("Contract Type",
     "Courtesy room block — 10 rooms per night; NO financial responsibility to client",
     "Attrition contract — client financially responsible for 90% of contracted room block; "
     "minimum 10 rooms per night required"),
    ("Distance to Biltmore Estate", "0.2 miles", "1.8 miles"),
    ("Reservation Cutoff",
     "30 days prior to arrival — by April 14, 2027",
     "60 days prior to arrival — by March 15, 2027"),
]

table = doc.add_table(rows=1 + len(rows), cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
widths = [Inches(1.5), Inches(2.6), Inches(2.6)]

hdr = table.rows[0].cells
for i, text in enumerate(["", "Grand Bohemian Hotel Asheville", "The Radical"]):
    hdr[i].text = ""
    p = hdr[i].paragraphs[0]
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = HEADER_TEXT
    set_cell_bg(hdr[i], ACCENT)

for ri, (label, c1, c2) in enumerate(rows, start=1):
    cells = table.rows[ri].cells
    p = cells[0].paragraphs[0]
    r = p.add_run(label); r.bold = True
    set_cell_bg(cells[0], LIGHT)
    for ci, content in ((1, c1), (2, c2)):
        cell = cells[ci]
        if content == "WEBSITE1":
            add_hyperlink(cell.paragraphs[0], "https://www.kesslercollection.com/bohemian-asheville/",
                          "kesslercollection.com/bohemian-asheville")
        elif content == "WEBSITE2":
            add_hyperlink(cell.paragraphs[0], "https://www.theradicalhotel.com", "theradicalhotel.com")
        elif "TALINK" in content:
            before = content.split("(TALINK)")[0]
            cell.paragraphs[0].add_run(before + "(")
            add_hyperlink(cell.paragraphs[0],
                          "https://www.tripadvisor.com/Hotel_Review-g60742-d1505353-Reviews-Grand_Bohemian_Hotel_Asheville_Autograph_Collection-Asheville_North_Carolina.html",
                          "TripAdvisor listing")
            cell.paragraphs[0].add_run(")")
        else:
            lines = content.split("\n")
            cell.paragraphs[0].add_run(lines[0])
            for line in lines[1:]:
                cell.add_paragraph(line)
    if ri % 2 == 0:
        for ci in (1, 2):
            set_cell_bg(cells[ci], "F7F4F9")

# set column widths
for row in table.rows:
    for idx, w in enumerate(widths):
        row.cells[idx].width = w

# ---------- Hotel 1 detail ----------
doc.add_page_break()
heading("Option 1: Grand Bohemian Hotel Asheville", size=15, space_before=0)
p = doc.add_paragraph()
p.add_run("11 Boston Way, Asheville, NC 28803  |  ").italic = True
add_hyperlink(p, "https://www.kesslercollection.com/bohemian-asheville/", "kesslercollection.com/bohemian-asheville")

heading("Ranking & Reputation", size=12, space_before=10)
p = doc.add_paragraph()
p.add_run("Ranked #18 of 28 hotels in Asheville on TripAdvisor (")
add_hyperlink(p, "https://www.tripadvisor.com/Hotel_Review-g60742-d1505353-Reviews-Grand_Bohemian_Hotel_Asheville_Autograph_Collection-Asheville_North_Carolina.html", "view listing")
p.add_run(").")

heading("Group Rates", size=12, space_before=10)
for item in [
    "Premier Two Queen Mountain View Suite with Balcony and Kitchenette — $389/night (regular rate $599)",
    "Classic King — $339/night (regular rate $529)",
]:
    doc.add_paragraph(item, style="List Bullet")

heading("Features & Amenities", size=12, space_before=10)
for item in [
    "On-site restaurant & bar — Red Stag Grill",
    "Art gallery",
    "Fitness center",
    "Valet parking: $35/night (subject to change)",
]:
    doc.add_paragraph(item, style="List Bullet")

heading("Concessions", size=12, space_before=10)
for item in [
    "Complimentary Wi-Fi",
    "Complimentary gift bag distribution at the front desk (non-personalized)",
]:
    doc.add_paragraph(item, style="List Bullet")

heading("Contract Terms", size=12, space_before=10)
p = doc.add_paragraph()
p.add_run("Courtesy room block: ").bold = True
p.add_run("The hotel can offer 10 rooms per night on a courtesy agreement. A courtesy "
          "agreement carries NO financial responsibility to the client — unbooked rooms are "
          "simply released at the cutoff date with no penalty.")

heading("Logistics & Deadlines", size=12, space_before=10)
for item in [
    "Distance to Biltmore Estate: 0.2 miles",
    "Reservation cutoff: 30 days prior to arrival — guests must book by April 14, 2027",
]:
    doc.add_paragraph(item, style="List Bullet")

# ---------- Hotel 2 detail ----------
doc.add_page_break()
heading("Option 2: The Radical", size=15, space_before=0)
p = doc.add_paragraph()
p.add_run("95 Roberts St, Asheville, NC 28801  |  ").italic = True
add_hyperlink(p, "https://www.theradicalhotel.com", "theradicalhotel.com")

heading("Ranking & Reputation", size=12, space_before=10)
doc.add_paragraph(
    "No TripAdvisor ranking was found for this property. The Radical is a newer hotel and does "
    "not yet have an established review ranking."
)

heading("Group Rates", size=12, space_before=10)
for item in [
    "King Studio — $279/night (regular rate $409)",
    "Double Queen Studio — $289/night (regular rate $419)",
]:
    doc.add_paragraph(item, style="List Bullet")

heading("Features & Amenities", size=12, space_before=10)
for item in [
    "Rooftop restaurant & bar — Golden Hour",
    "On-site café",
    "Fitness center",
    "Self-parking: $25/night (subject to change)",
]:
    doc.add_paragraph(item, style="List Bullet")

heading("Concessions", size=12, space_before=10)
for item in [
    "Complimentary Wi-Fi",
    "Complimentary welcome drink at check-in",
]:
    doc.add_paragraph(item, style="List Bullet")

heading("Contract Terms", size=12, space_before=10)
p = doc.add_paragraph()
p.add_run("90% attrition contract: ").bold = True
p.add_run("The client is financially responsible for 90% of the contracted room block, with a "
          "minimum of 10 rooms per night required. If guests do not book at least 90% of the "
          "contracted rooms, the client must pay for the shortfall.")

heading("Logistics & Deadlines", size=12, space_before=10)
for item in [
    "Distance to Biltmore Estate: 1.8 miles",
    "Reservation cutoff: 60 days prior to arrival — guests must book by March 15, 2027",
]:
    doc.add_paragraph(item, style="List Bullet")

# ---------- Key considerations ----------
heading("Key Considerations", size=15)
considerations = [
    ("Financial risk", "The Grand Bohemian's courtesy agreement carries no financial responsibility, "
     "while The Radical's 90% attrition clause puts the client on the hook for 90% of a minimum "
     "10-room-per-night block if guests do not book. This is the most significant difference between "
     "the two options."),
    ("Rate", "The Radical's group rates are $50–$110 per night lower than the Grand Bohemian's, "
     "which may matter to budget-conscious guests."),
    ("Location", "The Grand Bohemian is essentially at the Biltmore Estate's doorstep (0.2 miles) "
     "versus 1.8 miles for The Radical — convenient for guests and helpful for shuttle logistics."),
    ("Booking window", "The Radical's cutoff (March 15, 2027; 60 days out) is a full month earlier "
     "than the Grand Bohemian's (April 14, 2027; 30 days out), leaving less time for late-booking guests."),
    ("Reputation", "The Grand Bohemian has an established TripAdvisor presence (#18 of 28 in "
     "Asheville); The Radical is newer and has no ranking yet, so guest-experience data is limited."),
]
for label, text in considerations:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(label + ": ").bold = True
    p.add_run(text)

heading("Recommendation", size=15)
doc.add_paragraph(
    "The Grand Bohemian Hotel Asheville is the lower-risk choice: its courtesy agreement carries no "
    "financial exposure, it sits just 0.2 miles from the Biltmore Estate, and its later cutoff date "
    "(April 14, 2027) gives guests more time to book. The Radical offers meaningfully lower nightly "
    "rates and a newer, design-forward property, but the 90% attrition clause on a 10-room-per-night "
    "minimum creates real financial risk if bookings fall short, and the earlier March 15, 2027 "
    "cutoff compresses the booking window. Unless guest budget is the overriding concern and strong "
    "pickup is certain, the Grand Bohemian's courtesy block is recommended."
)

doc.save(OUT)
print("Saved:", OUT)
